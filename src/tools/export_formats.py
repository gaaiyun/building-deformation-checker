"""共享的报告导出格式：DOCX 与 HTML

把原本嵌在 app.py 内的 _generate_docx / _generate_html 抽出，
让 Streamlit、PySide6 桌面、CLI 等不同 UI 共用同一套导出逻辑。
"""

from __future__ import annotations

from contextlib import contextmanager
import io
import os
from pathlib import Path
import tempfile
from datetime import datetime

from src.tools.extraction_quality import append_issue_source_hint


def generate_docx(md_content: str, report, errors: list, warnings: list) -> bytes:
    """生成 Word 文档（.docx 二进制内容）"""
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()

    def set_run_style(run, *, size: float = 10.5, bold: bool = False) -> None:
        run.font.name = "Microsoft YaHei"
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    def style_paragraph(
        paragraph,
        *,
        size: float = 10.5,
        bold: bool = False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before: float = 0,
        space_after: float = 4,
        line_spacing: float = 1.35,
    ) -> None:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)
        paragraph.paragraph_format.line_spacing = line_spacing
        if not paragraph.runs:
            run = paragraph.add_run("")
            set_run_style(run, size=size, bold=bold)
        for run in paragraph.runs:
            set_run_style(run, size=size, bold=bool(run.bold) or bold)

    def shade_cell(cell, fill: str = "F3F4F6") -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, fill: str | None = None) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text if text else "-")
        set_run_style(run, size=10.5, bold=bold)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if fill:
            shade_cell(cell, fill)

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_rpr = normal_style._element.get_or_add_rPr()
    normal_rfonts = normal_rpr.rFonts
    if normal_rfonts is None:
        normal_rfonts = OxmlElement("w:rFonts")
        normal_rpr.append(normal_rfonts)
    normal_rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    title = doc.add_paragraph()
    title_run = title.add_run("建筑变形监测报告检查报告")
    set_run_style(title_run, size=18, bold=True)
    style_paragraph(title, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("自动核验输出文件")
    set_run_style(subtitle_run, size=10.5)
    style_paragraph(subtitle, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    overview_heading = doc.add_paragraph()
    overview_run = overview_heading.add_run("一、报告概览")
    set_run_style(overview_run, size=13, bold=True)
    style_paragraph(overview_heading, size=13, bold=True, space_after=6)

    overview = doc.add_table(rows=5, cols=2, style="Table Grid")
    overview.alignment = WD_TABLE_ALIGNMENT.CENTER
    overview_rows = [
        ("项目名称", report.project_name or "-"),
        ("监测单位", report.monitoring_company or "-"),
        ("报告编号", report.report_number or "-"),
        ("监测日期", report.monitoring_date or "-"),
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    ]
    for row_idx, (label, value) in enumerate(overview_rows):
        set_cell_text(overview.cell(row_idx, 0), label, bold=True, fill="F3F4F6")
        set_cell_text(overview.cell(row_idx, 1), str(value))

    doc.add_paragraph("")

    diagnostics = report.extraction_diagnostics or {}
    extraction_label = diagnostics.get("method", "unknown")
    if diagnostics.get("selected_profile"):
        extraction_label += f" / {diagnostics['selected_profile']}"

    summary_heading = doc.add_paragraph()
    summary_run = summary_heading.add_run("二、检查摘要")
    set_run_style(summary_run, size=13, bold=True)
    style_paragraph(summary_heading, size=13, bold=True, space_after=6)

    summary = doc.add_table(rows=5, cols=4, style="Table Grid")
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(["指标", "数值", "指标", "数值"]):
        set_cell_text(summary.cell(0, idx), header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="EDEFF3")

    summary_pairs = [
        ("数据表", str(len(report.tables)), "错误", str(len(errors))),
        ("警告", str(len(warnings)), "提取方式", extraction_label),
        ("原始字符", f"{diagnostics.get('raw_chars', 0):,}", "清洗后字符", f"{diagnostics.get('clean_chars', 0):,}"),
        ("压缩率", f"{diagnostics.get('compression_ratio', 0.0):.1%}", "疑似异常表", str(diagnostics.get("abnormal_table_count", 0))),
    ]
    for row_idx, values in enumerate(summary_pairs, start=1):
        for col_idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx % 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(summary.cell(row_idx, col_idx), value, align=align)

    if report.table_extraction_flags:
        diag_note = doc.add_paragraph()
        diag_run = diag_note.add_run(
            "提取提示："
            + "；".join(
                f"{report.tables[idx].monitoring_item}{f' ({report.tables[idx].borehole_id})' if report.tables[idx].borehole_id else ''}：{'；'.join(flags)}"
                for idx, flags in sorted(report.table_extraction_flags.items())
                if idx < len(report.tables)
            )
        )
        set_run_style(diag_run, size=10)
        style_paragraph(diag_note, size=10, space_before=4, space_after=6)

    doc.add_paragraph("")

    issues_heading = doc.add_paragraph()
    issues_run = issues_heading.add_run("三、问题清单")
    set_run_style(issues_run, size=13, bold=True)
    style_paragraph(issues_heading, size=13, bold=True, space_after=6)

    def add_issue_table(title_text: str, issue_list: list) -> None:
        title_p = doc.add_paragraph()
        title_r = title_p.add_run(title_text)
        set_run_style(title_r, size=11.5, bold=True)
        style_paragraph(title_p, size=11.5, bold=True, space_before=2, space_after=4)
        if not issue_list:
            empty_p = doc.add_paragraph()
            empty_r = empty_p.add_run("未发现需关注的问题。")
            set_run_style(empty_r, size=10.5)
            style_paragraph(empty_p, size=10.5, space_after=6)
            return
        table = doc.add_table(rows=1, cols=5, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["序号", "表名", "测点", "字段", "说明"]
        for idx, header in enumerate(headers):
            set_cell_text(table.cell(0, idx), header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="EDEFF3")
        for idx, issue in enumerate(issue_list, start=1):
            row = table.add_row().cells
            set_cell_text(row[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row[1], issue.table_name)
            set_cell_text(row[2], issue.point_id or "-", align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row[3], issue.field_name or "-", align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row[4], append_issue_source_hint(issue.message, issue.suspected_source))
        doc.add_paragraph("")

    add_issue_table("3.1 错误项", errors)
    add_issue_table("3.2 警告项", warnings)

    conclusion_heading = doc.add_paragraph()
    conclusion_run = conclusion_heading.add_run("四、结论")
    set_run_style(conclusion_run, size=13, bold=True)
    style_paragraph(conclusion_heading, size=13, bold=True, space_after=6)

    if report.conclusion:
        original_conclusion = doc.add_paragraph()
        original_run = original_conclusion.add_run(f"报告原文结论：{report.conclusion}")
        set_run_style(original_run, size=10.5)
        style_paragraph(original_conclusion, size=10.5, space_after=4)

    system_conclusion = doc.add_paragraph()
    if errors:
        result_text = f"自动检查结论：共发现 {len(errors)} 条错误、{len(warnings)} 条警告，建议结合原始报告进行人工复核。"
    elif warnings:
        result_text = f"自动检查结论：未发现错误，存在 {len(warnings)} 条警告，建议按需复核。"
    else:
        result_text = "自动检查结论：本次自动检查未发现错误或警告。"
    result_run = system_conclusion.add_run(result_text)
    set_run_style(result_run, size=10.5)
    style_paragraph(system_conclusion, size=10.5, space_after=10)

    footer = doc.add_paragraph()
    footer_run = footer.add_run("本文件由建筑变形监测报告核验台自动生成。")
    set_run_style(footer_run, size=9)
    style_paragraph(footer, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_html(md_content: str, project_name: str) -> str:
    """生成可打印的 HTML 报告（可在浏览器中打印为 PDF）"""
    import markdown

    html_body = markdown.markdown(
        md_content,
        extensions=["tables", "fenced_code"],
    )

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>{project_name} - 检查报告</title>
    <style>
        body {{ font-family: "Microsoft YaHei", "SimHei", sans-serif; max-width: 900px; margin: 0 auto; padding: 20px; color: #333; }}
        h1 {{ color: #1a5276; border-bottom: 3px solid #1a5276; padding-bottom: 10px; }}
        h2 {{ color: #2c3e50; border-bottom: 1px solid #bdc3c7; padding-bottom: 6px; margin-top: 30px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
        th {{ background-color: #f2f2f2; font-weight: bold; }}
        tr:nth-child(even) {{ background-color: #fafafa; }}
        blockquote {{ border-left: 4px solid #3498db; margin: 15px 0; padding: 10px 20px; background: #ecf6fd; }}
        code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-size: 0.9em; }}
        hr {{ border: none; border-top: 1px solid #ddd; margin: 30px 0; }}
        @media print {{
            body {{ max-width: 100%; padding: 0; }}
            h1 {{ page-break-before: avoid; }}
            table {{ page-break-inside: avoid; }}
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""


def _enum_value(value) -> str:
    return getattr(value, "value", str(value or ""))


def _fmt(value) -> str | float | int:
    if value is None:
        return ""
    if isinstance(value, (int, float)):
        return value
    return str(value)


def _is_xml_compatible_char(ch: str) -> bool:
    codepoint = ord(ch)
    return (
        codepoint in (0x09, 0x0A, 0x0D)
        or 0x20 <= codepoint <= 0xD7FF
        or 0xE000 <= codepoint <= 0xFFFD
        or 0x10000 <= codepoint <= 0x10FFFF
    )


def _sanitize_excel_text(value: str) -> str:
    """Replace characters that cannot be serialized into XLSX XML parts."""
    if all(_is_xml_compatible_char(ch) for ch in value):
        return value
    return "".join(ch if _is_xml_compatible_char(ch) else "\uFFFD" for ch in value)


def _default_excel_temp_dir() -> Path | None:
    configured = os.getenv("BDC_EXCEL_TEMP_DIR")
    if configured:
        return Path(configured)

    g_cache = Path("G:/dev-cache/building-deformation-checker/openpyxl-temp")
    if g_cache.drive and Path("G:/").exists():
        return g_cache

    return None


@contextmanager
def _openpyxl_tempdir():
    """Route openpyxl worksheet temp files away from a full system temp drive."""
    temp_dir = _default_excel_temp_dir()
    if temp_dir is None:
        yield
        return

    temp_dir.mkdir(parents=True, exist_ok=True)
    old_tempdir = tempfile.tempdir
    old_tmp = os.environ.get("TMP")
    old_temp = os.environ.get("TEMP")
    tempfile.tempdir = str(temp_dir)
    os.environ["TMP"] = str(temp_dir)
    os.environ["TEMP"] = str(temp_dir)
    try:
        yield
    finally:
        tempfile.tempdir = old_tempdir
        if old_tmp is None:
            os.environ.pop("TMP", None)
        else:
            os.environ["TMP"] = old_tmp
        if old_temp is None:
            os.environ.pop("TEMP", None)
        else:
            os.environ["TEMP"] = old_temp


def _safe_excel_value(value) -> str | float | int:
    """避免 PDF/LLM 文本被 Excel 当公式执行。

    数值型负数会以 int/float 写入，不受影响；只有字符串字段才会在疑似公式
    前缀前加引号。
    """
    value = _fmt(value)
    if isinstance(value, str):
        value = _sanitize_excel_text(value)
        stripped = value.lstrip()
        if stripped.startswith(("=", "+", "@")):
            return "'" + value
        if stripped.startswith("-") and stripped[1:2].isalpha():
            return "'" + value
    return value


def _table_name(table) -> str:
    name = getattr(table, "monitoring_item", "") or "未命名表"
    borehole_id = getattr(table, "borehole_id", "")
    if borehole_id:
        return f"{name} ({borehole_id})"
    return name


def generate_intermediate_xlsx(
    report,
    *,
    calc_issues: list | None = None,
    stats_issues: list | None = None,
    logic_issues: list | None = None,
    analysis_plan: list[dict] | None = None,
) -> bytes:
    """生成可审查的 Excel 中间层。

    这个工作簿不是最终 Word 报告的替代品，而是把已经进入规则引擎的结构化数据
    展开给业务人员核查：表格清单、标准化测点、深层位移、统计摘要、问题清单和
    分析计划。后续若新增 LLM 前候选表格，也应继续复用这个导出入口。
    """
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    calc_issues = calc_issues or []
    stats_issues = stats_issues or []
    logic_issues = logic_issues or []
    analysis_plan = analysis_plan or []
    tables = getattr(report, "tables", []) if report is not None else []

    wb = Workbook()
    wb.remove(wb.active)

    header_fill = PatternFill("solid", fgColor="EAF1FB")
    section_fill = PatternFill("solid", fgColor="F7F9FC")
    header_font = Font(name="Microsoft YaHei", bold=True, color="0F2F5F")
    normal_font = Font(name="Microsoft YaHei", color="111827")
    note_font = Font(name="Microsoft YaHei", color="4B5563")
    wrap = Alignment(vertical="top", wrap_text=True)

    def add_sheet(title: str, headers: list[str]):
        ws = wb.create_sheet(title)
        ws.append(headers)
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        return ws

    def append_safe(ws, values: list) -> None:
        ws.append([_safe_excel_value(value) for value in values])

    def finish_sheet(
        ws,
        *,
        widths: dict[int, int] | None = None,
        style_body: bool = True,
        auto_width: bool = True,
    ) -> None:
        widths = widths or {}
        ws.auto_filter.ref = ws.dimensions
        if style_body:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.row == 1:
                        cell.font = header_font
                    else:
                        cell.font = normal_font
                        cell.alignment = wrap
        for col_idx in range(1, ws.max_column + 1):
            max_len = 8
            if auto_width:
                for column_cells in ws.iter_cols(min_col=col_idx, max_col=col_idx, min_row=1, max_row=ws.max_row):
                    for cell in column_cells:
                        value = cell.value
                        if value is not None:
                            max_len = max(max_len, min(len(str(value)) + 2, 40))
            ws.column_dimensions[get_column_letter(col_idx)].width = widths.get(col_idx, max_len)

    overview = wb.create_sheet("00_报告概览")
    overview.append(["字段", "值"])
    for cell in overview[1]:
        cell.fill = header_fill
        cell.font = header_font
    diagnostics = getattr(report, "extraction_diagnostics", {}) if report is not None else {}
    overview_rows = [
        ("项目名称", getattr(report, "project_name", "") if report is not None else ""),
        ("监测单位", getattr(report, "monitoring_company", "") if report is not None else ""),
        ("报告编号", getattr(report, "report_number", "") if report is not None else ""),
        ("监测日期", getattr(report, "monitoring_date", "") if report is not None else ""),
        ("监测时间段", getattr(report, "monitoring_period", "") if report is not None else ""),
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("结构化表数量", len(tables)),
        ("阈值数量", len(getattr(report, "thresholds", []) if report is not None else [])),
        ("汇总项数量", len(getattr(report, "summary_items", []) if report is not None else [])),
        ("计算问题数量", len(calc_issues)),
        ("统计问题数量", len(stats_issues)),
        ("逻辑问题数量", len(logic_issues)),
        ("提取方式", diagnostics.get("method", "")),
        ("OCR/清洗配置", diagnostics.get("selected_profile", "")),
        ("原始字符", diagnostics.get("raw_chars", "")),
        ("清洗后字符", diagnostics.get("clean_chars", "")),
        ("压缩率", diagnostics.get("compression_ratio", "")),
        ("异常表数量", diagnostics.get("abnormal_table_count", "")),
        ("原始候选表数量", diagnostics.get("raw_table_candidate_count", len(diagnostics.get("raw_table_candidates", []) or []))),
    ]
    for label, value in overview_rows:
        append_safe(overview, [label, value])
    for row in overview.iter_rows():
        for cell in row:
            cell.font = header_font if cell.column == 1 else normal_font
            cell.alignment = wrap
            if cell.column == 1 and cell.row > 1:
                cell.fill = section_fill
    overview.column_dimensions["A"].width = 18
    overview.column_dimensions["B"].width = 48

    raw_candidates = diagnostics.get("raw_table_candidates", []) or []
    candidate_ws = add_sheet(
        "00A_候选表清单",
        ["候选表ID", "引擎", "来源页", "页内序号", "行数", "列数", "标题/首行预览", "质量标志"],
    )
    for candidate in raw_candidates:
        flags = candidate.get("quality_flags", []) or []
        append_safe(candidate_ws, [
            candidate.get("table_id", ""),
            candidate.get("engine", ""),
            candidate.get("page", ""),
            candidate.get("table_index", ""),
            candidate.get("row_count", ""),
            candidate.get("col_count", ""),
            candidate.get("title_preview", ""),
            "；".join(str(flag) for flag in flags),
        ])
    finish_sheet(candidate_ws, widths={1: 16, 7: 72, 8: 32})

    raw_cell_headers = ["候选表ID", "引擎", "来源页", "原始行号", "原始列号", "原始单元格值"]
    raw_cell_ws = add_sheet("00B_候选表单元格", raw_cell_headers)
    raw_cell_sheet_index = 1
    raw_cell_row_count = 1
    max_raw_cell_rows = 900_000
    for candidate in raw_candidates:
        for row_index, row in enumerate(candidate.get("rows", []) or [], start=1):
            for column_index, value in enumerate(row or [], start=1):
                if raw_cell_row_count >= max_raw_cell_rows:
                    finish_sheet(
                        raw_cell_ws,
                        widths={1: 16, 2: 14, 3: 10, 4: 12, 5: 12, 6: 48},
                        style_body=False,
                        auto_width=False,
                    )
                    raw_cell_sheet_index += 1
                    raw_cell_ws = add_sheet(
                        f"00B_候选表单元格_{raw_cell_sheet_index}",
                        raw_cell_headers,
                    )
                    raw_cell_row_count = 1
                append_safe(raw_cell_ws, [
                    candidate.get("table_id", ""),
                    candidate.get("engine", ""),
                    candidate.get("page", ""),
                    row_index,
                    column_index,
                    value,
                ])
                raw_cell_row_count += 1
    finish_sheet(
        raw_cell_ws,
        widths={1: 16, 2: 14, 3: 10, 4: 12, 5: 12, 6: 48},
        style_body=False,
        auto_width=False,
    )

    table_ws = add_sheet(
        "01_表格清单",
        [
            "序号",
            "表名",
            "类别",
            "日期",
            "期次",
            "测孔",
            "声明测点数",
            "解析测点数",
            "深层点数",
            "单位",
            "单位换算",
            "间隔天数",
            "初始值可靠",
            "提取提示",
            "来源分块",
            "来源页",
        ],
    )
    for idx, table in enumerate(tables, start=1):
        cfg = getattr(table, "verification_config", None)
        flags = getattr(report, "table_extraction_flags", {}).get(idx - 1, []) if report is not None else []
        append_safe(table_ws, [
            idx,
            _table_name(table),
            _enum_value(getattr(table, "category", "")),
            _fmt(getattr(table, "monitor_date", "")),
            _fmt(getattr(table, "monitor_count", "")),
            _fmt(getattr(table, "borehole_id", "")),
            _fmt(getattr(table, "point_count", "")),
            len(getattr(table, "points", []) or []),
            len(getattr(table, "deep_points", []) or []),
            _fmt(getattr(cfg, "unit", "")),
            _fmt(getattr(cfg, "unit_conversion", "")),
            _fmt(getattr(cfg, "interval_days", "")),
            _fmt(getattr(cfg, "initial_value_reliable", "")),
            "；".join(flags),
            _fmt(getattr(table, "source_chunk", 0)),
            _fmt(getattr(table, "source_pages", "")),
        ])
    finish_sheet(table_ws, widths={2: 28, 14: 42, 16: 14})

    point_ws = add_sheet(
        "02_标准化测点",
        [
            "表序号",
            "表名",
            "类别",
            "日期",
            "测点",
            "初始值",
            "上次值",
            "本次值",
            "本次变化",
            "累计变化",
            "变化速率",
            "安全状态",
            "来源分块",
            "来源页",
            "原始行",
            "字段列映射",
        ],
    )
    for idx, table in enumerate(tables, start=1):
        for point in getattr(table, "points", []) or []:
            append_safe(point_ws, [
                idx,
                _table_name(table),
                _enum_value(getattr(table, "category", "")),
                _fmt(getattr(table, "monitor_date", "")),
                _fmt(getattr(point, "point_id", "")),
                _fmt(getattr(point, "initial_value", None)),
                _fmt(getattr(point, "previous_value", None)),
                _fmt(getattr(point, "current_value", None)),
                _fmt(getattr(point, "current_change", None)),
                _fmt(getattr(point, "cumulative_change", None)),
                _fmt(getattr(point, "change_rate", None)),
                _fmt(getattr(point, "safety_status", "")),
                _fmt(getattr(point, "source_chunk", 0)),
                _fmt(getattr(point, "source_page", None)),
                _fmt(getattr(point, "source_row_text", "")),
                _fmt(getattr(point, "source_field_map", "")),
            ])
    finish_sheet(point_ws, widths={2: 28, 15: 72, 16: 44})

    deep_ws = add_sheet(
        "03_深层位移",
        [
            "表序号",
            "表名",
            "日期",
            "测孔",
            "深度",
            "上次累计",
            "本次累计",
            "本次变化",
            "变化速率",
            "来源分块",
            "来源页",
            "原始行",
            "字段列映射",
        ],
    )
    for idx, table in enumerate(tables, start=1):
        for point in getattr(table, "deep_points", []) or []:
            append_safe(deep_ws, [
                idx,
                _table_name(table),
                _fmt(getattr(table, "monitor_date", "")),
                _fmt(getattr(table, "borehole_id", "")),
                _fmt(getattr(point, "depth", None)),
                _fmt(getattr(point, "previous_cumulative", None)),
                _fmt(getattr(point, "current_cumulative", None)),
                _fmt(getattr(point, "current_change", None)),
                _fmt(getattr(point, "change_rate", None)),
                _fmt(getattr(point, "source_chunk", 0)),
                _fmt(getattr(point, "source_page", None)),
                _fmt(getattr(point, "source_row_text", "")),
                _fmt(getattr(point, "source_field_map", "")),
            ])
    finish_sheet(deep_ws, widths={2: 28, 12: 72, 13: 44})

    stats_ws = add_sheet(
        "04_统计摘要",
        [
            "表序号",
            "表名",
            "正向最大点",
            "正向最大值",
            "负向最大点",
            "负向最大值",
            "最大速率点",
            "最大速率值",
            "最大变化点",
            "最大变化值",
            "最大力值点",
            "最大力值",
            "最小力值点",
            "最小力值",
        ],
    )
    for idx, table in enumerate(tables, start=1):
        stats = getattr(table, "statistics", None)
        append_safe(stats_ws, [
            idx,
            _table_name(table),
            _fmt(getattr(stats, "positive_max_id", "")),
            _fmt(getattr(stats, "positive_max_value", None)),
            _fmt(getattr(stats, "negative_max_id", "")),
            _fmt(getattr(stats, "negative_max_value", None)),
            _fmt(getattr(stats, "max_rate_id", "")),
            _fmt(getattr(stats, "max_rate_value", None)),
            _fmt(getattr(stats, "max_change_id", "")),
            _fmt(getattr(stats, "max_change_value", None)),
            _fmt(getattr(stats, "max_force_id", "")),
            _fmt(getattr(stats, "max_force_value", None)),
            _fmt(getattr(stats, "min_force_id", "")),
            _fmt(getattr(stats, "min_force_value", None)),
        ])
    finish_sheet(stats_ws, widths={2: 28})

    issue_ws = add_sheet(
        "05_问题清单",
        ["来源", "级别", "表名", "测点", "字段", "期望", "实际", "说明"],
    )
    for source_name, issues in (
        ("计算核验", calc_issues),
        ("统计核验", stats_issues),
        ("逻辑检查", logic_issues),
    ):
        for issue in issues:
            append_safe(issue_ws, [
                source_name,
                _fmt(getattr(issue, "severity", "")),
                _fmt(getattr(issue, "table_name", "")),
                _fmt(getattr(issue, "point_id", "")),
                _fmt(getattr(issue, "field_name", "")),
                _fmt(getattr(issue, "expected_value", "")),
                _fmt(getattr(issue, "actual_value", "")),
                append_issue_source_hint(
                    _fmt(getattr(issue, "message", "")),
                    _fmt(getattr(issue, "suspected_source", "")),
                ),
            ])
    finish_sheet(issue_ws, widths={3: 28, 8: 60})

    plan_ws = add_sheet(
        "06_分析计划",
        ["表序号", "表名", "类别", "测点数", "单位", "单位换算", "间隔天数", "计算方法", "特殊说明"],
    )
    for plan in analysis_plan:
        methods = []
        for method in plan.get("verification_methods", []) or []:
            if isinstance(method, dict):
                methods.append(
                    f"{method.get('name', '')}: {method.get('formula', '')}"
                    + (f" (tol={method.get('tolerance')})" if method.get("tolerance") else "")
                )
            else:
                methods.append(str(method))
        notes = plan.get("special_notes", []) or []
        append_safe(plan_ws, [
            _fmt(plan.get("table_index", "")),
            _fmt(plan.get("table_name", "")),
            _fmt(plan.get("category", "")),
            _fmt(plan.get("point_count", "")),
            _fmt(plan.get("unit", "")),
            _fmt(plan.get("unit_conversion", "")),
            _fmt(plan.get("interval_days", "")),
            "\n".join(methods),
            "\n".join(str(n) for n in notes),
        ])
    finish_sheet(plan_ws, widths={2: 28, 8: 60, 9: 48})

    for ws in wb.worksheets:
        ws.sheet_view.showGridLines = False
        if not ws.title.startswith("00B_候选表单元格"):
            for row in ws.iter_rows(min_row=2):
                if row and row[0].row % 2 == 0:
                    for cell in row:
                        cell.fill = section_fill
        if ws.max_row == 1:
            ws.append(["", "无数据"])
            ws["B2"].font = note_font

    buf = io.BytesIO()
    with _openpyxl_tempdir():
        wb.save(buf)
    return buf.getvalue()


__all__ = ["generate_docx", "generate_html", "generate_intermediate_xlsx"]
