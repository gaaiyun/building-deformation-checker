"""
建筑变形监测报告检查智能体 — Streamlit Web UI

功能:
- 上传PDF监测报告，自动提取数据并检查
- 实时进度与日志显示
- 多格式导出（Markdown / Word / HTML）
- 分类展示检查结果
"""

import io
import logging
import tempfile
import time
from collections import defaultdict
from datetime import datetime
from pathlib import Path

import streamlit as st
from src.tools.extraction_quality import append_issue_source_hint

# ── 日志配置：捕获到 StreamHandler 供界面显示 ──────────
log_records: list[str] = []


class StreamlitLogHandler(logging.Handler):
    """把日志记录收集到列表，供 UI 实时展示"""
    def emit(self, record):
        log_records.append(self.format(record))


handler = StreamlitLogHandler()
handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(name)s: %(message)s", datefmt="%H:%M:%S"))
root_logger = logging.getLogger()
root_logger.setLevel(logging.INFO)
root_logger.addHandler(handler)
logger = logging.getLogger("app")


# ── 辅助函数（必须在 Streamlit UI 逻辑之前定义）─────────

def _render_issues(title: str, issues: list) -> None:
    """按表名分组展示检查问题"""
    if not issues:
        st.success(f"{title}全部通过")
        return

    errors = [i for i in issues if i.severity == "error"]
    warnings = [i for i in issues if i.severity == "warning"]
    infos = [i for i in issues if i.severity == "info"]

    if errors:
        st.error(f"发现 {len(errors)} 个错误")
    if warnings:
        st.warning(f"发现 {len(warnings)} 个警告")
    if infos:
        st.info(f"发现 {len(infos)} 个提示")

    grouped = defaultdict(list)
    for issue in issues:
        grouped[issue.table_name].append(issue)

    for table_name, table_issues in grouped.items():
        err_count = sum(1 for i in table_issues if i.severity == "error")
        warn_count = sum(1 for i in table_issues if i.severity == "warning")
        badge_parts = []
        if err_count:
            badge_parts.append(f"E{err_count}")
        if warn_count:
            badge_parts.append(f"W{warn_count}")
        badge = f"[{' / '.join(badge_parts)}]" if badge_parts else ""

        with st.expander(f"**{table_name}** {badge}", expanded=bool(err_count)):
            for issue in table_issues:
                message = append_issue_source_hint(issue.message, issue.suspected_source)
                if issue.severity == "error":
                    st.error(f"**{issue.point_id}** | {issue.field_name}: {message}")
                elif issue.severity == "warning":
                    st.warning(f"**{issue.point_id}** | {issue.field_name}: {message}")
                else:
                    st.info(message)


def _render_extraction_diagnostics(report) -> None:
    diagnostics = report.extraction_diagnostics or {}
    if not diagnostics:
        return

    st.markdown("### 提取诊断")
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("原始字符", f"{diagnostics.get('raw_chars', 0):,}")
    c2.metric("清洗后字符", f"{diagnostics.get('clean_chars', 0):,}")
    ratio = diagnostics.get("compression_ratio", 0.0)
    c3.metric("压缩率", f"{ratio:.1%}")
    c4.metric("异常页", f"{len(diagnostics.get('high_markup_pages', []))}")
    c5.metric("异常表", f"{diagnostics.get('abnormal_table_count', 0)}")

    method = diagnostics.get("method", "unknown")
    profile = diagnostics.get("selected_profile", "")
    label = f"{method} ({profile})" if profile else method
    st.caption(f"提取方式: {label}")

    attempts = diagnostics.get("attempts", [])
    if attempts:
        st.markdown("**提取尝试链路**")
        for attempt in attempts:
            if attempt.get("error"):
                st.warning(f"{attempt['profile']}: {attempt['error']}")
            else:
                st.caption(
                    f"{attempt['profile']}: clean={attempt.get('clean_chars', 0):,} chars, "
                    f"pages={attempt.get('page_count', 0)}, "
                    f"compression={attempt.get('compression_ratio', 0.0):.1%}"
                )

    high_markup_pages = diagnostics.get("high_markup_pages", [])
    if high_markup_pages:
        page_label = "，".join(str(page + 1) for page in high_markup_pages[:12])
        suffix = " ..." if len(high_markup_pages) > 12 else ""
        st.caption(f"高 markup 页: 第 {page_label} 页{suffix}")

    debug_dir = diagnostics.get("debug_dir", "")
    if debug_dir:
        st.markdown("**OCR 调试目录**")
        st.code(debug_dir, language=None)

    flagged_tables = report.table_extraction_flags or {}
    if flagged_tables:
        st.markdown("**疑似提取异常表**")
        for table_index, flags in sorted(flagged_tables.items()):
            if table_index >= len(report.tables):
                continue
            table = report.tables[table_index]
            table_name = table.monitoring_item
            if table.borehole_id:
                table_name += f" ({table.borehole_id})"
            st.warning(f"{table_name}: {'；'.join(flags)}")


def _render_analysis_plan(analysis_plan: list[dict]) -> None:
    """ReAct 风格渲染每张表的分析计划：Thought → Observation → Action"""
    if not analysis_plan:
        st.info("未生成分析计划")
        return

    FIELD_LABELS = {
        "initial_value": "初始值",
        "previous_value": "上次值",
        "current_value": "本次值",
        "current_change": "本次变化",
        "cumulative_change": "累计变化",
        "change_rate": "速率",
        "safety_status": "安全状态",
        "depth": "深度",
        "previous_cumulative": "上次累计",
        "current_cumulative": "本次累计",
    }

    for plan in analysis_plan:
        has_notes = bool(plan["special_notes"])
        title_badge = " [需关注]" if has_notes else ""
        expander_title = (
            f"Table {plan['table_index']}: {plan['table_name']} "
            f"({plan['category']} | {plan['point_count']}个测点){title_badge}"
        )

        with st.expander(expander_title, expanded=has_notes):
            # ── Thought: 字段识别 ──────────────────────
            st.markdown("**字段识别**")
            fields = plan["fields_detected"]
            if fields:
                cols = st.columns(len(fields))
                for i, (field_key, detected) in enumerate(fields.items()):
                    label = FIELD_LABELS.get(field_key, field_key)
                    status = "是" if detected else "否"
                    cols[i].markdown(f"<div style='text-align:center'><b>{label}</b><br/>{status}</div>", unsafe_allow_html=True)
            st.markdown("")

            # ── Observation: 数据样本 ──────────────────
            st.markdown("**数据样本**")
            for sample in plan["data_sample"]:
                st.code(sample, language=None)

            # ── Observation: 单位与基准分析 ─────────────
            st.markdown("**单位与基准分析**")
            unit_text = f"单位: **{plan['unit']}**"
            if plan["unit_conversion"] != 1.0:
                unit_text += f" → mm (×{plan['unit_conversion']:.0f}转换)"
            else:
                unit_text += f" ({plan['conversion_note']})"
            st.markdown(unit_text)

            reliable_text = "可靠" if plan["initial_reliable"] else "需谨慎"
            st.markdown(f"初始值: {reliable_text}，{plan['reliability_reason']}")

            if plan["interval_days"]:
                st.markdown(f"监测间隔: **{plan['interval_days']:.0f}天** ({plan['interval_source']})")
            else:
                st.markdown(f"监测间隔: {plan['interval_source']}")
            st.markdown("")

            # ── Action: 验证规则 ────────────────────────
            st.markdown("**将执行的验证规则**")
            for method in plan["verification_methods"]:
                st.markdown(
                    f"**{method['name']}** = `{method['formula']}`, "
                    f"容差={method['tolerance']}, 级别={method['severity']}"
                )

            # ── 特殊说明 ────────────────────────────────
            if plan["special_notes"]:
                st.markdown("")
                st.markdown("**特殊说明**")
                for note in plan["special_notes"]:
                    st.warning(note)


def _generate_docx(md_content: str, report, errors: list, warnings: list) -> bytes:
    """生成 Word 文档"""
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()

    def set_run_style(run, *, size: float = 10.5, bold: bool = False) -> None:
        run.font.name = "Microsoft YaHei"
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    def style_paragraph(
        paragraph,
        *,
        size: float = 10.5,
        bold: bool = False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before: float = 0,
        space_after: float = 4,
        line_spacing: float = 1.35,
    ) -> None:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)
        paragraph.paragraph_format.line_spacing = line_spacing
        if not paragraph.runs:
            run = paragraph.add_run("")
            set_run_style(run, size=size, bold=bold)
        for run in paragraph.runs:
            set_run_style(run, size=size, bold=bool(run.bold) or bold)

    def shade_cell(cell, fill: str = "F3F4F6") -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, fill: str | None = None) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text if text else "-")
        set_run_style(run, size=10.5, bold=bold)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if fill:
            shade_cell(cell, fill)

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_rpr = normal_style._element.get_or_add_rPr()
    normal_rfonts = normal_rpr.rFonts
    if normal_rfonts is None:
        normal_rfonts = OxmlElement("w:rFonts")
        normal_rpr.append(normal_rfonts)
    normal_rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    title = doc.add_paragraph()
    title_run = title.add_run("建筑变形监测报告检查报告")
    set_run_style(title_run, size=18, bold=True)
    style_paragraph(title, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("自动核验输出文件")
    set_run_style(subtitle_run, size=10.5)
    style_paragraph(subtitle, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    overview_heading = doc.add_paragraph()
    overview_run = overview_heading.add_run("一、报告概览")
    set_run_style(overview_run, size=13, bold=True)
    style_paragraph(overview_heading, size=13, bold=True, space_after=6)

    overview = doc.add_table(rows=5, cols=2, style="Table Grid")
    overview.alignment = WD_TABLE_ALIGNMENT.CENTER
    overview_rows = [
        ("项目名称", report.project_name or "-"),
        ("监测单位", report.monitoring_company or "-"),
        ("报告编号", report.report_number or "-"),
        ("监测日期", report.monitoring_date or "-"),
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    ]
    for row_idx, (label, value) in enumerate(overview_rows):
        set_cell_text(overview.cell(row_idx, 0), label, bold=True, fill="F3F4F6")
        set_cell_text(overview.cell(row_idx, 1), str(value))

    doc.add_paragraph("")

    diagnostics = report.extraction_diagnostics or {}
    extraction_label = diagnostics.get("method", "unknown")
    if diagnostics.get("selected_profile"):
        extraction_label += f" / {diagnostics['selected_profile']}"

    summary_heading = doc.add_paragraph()
    summary_run = summary_heading.add_run("二、检查摘要")
    set_run_style(summary_run, size=13, bold=True)
    style_paragraph(summary_heading, size=13, bold=True, space_after=6)

    summary = doc.add_table(rows=5, cols=4, style="Table Grid")
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(["指标", "数值", "指标", "数值"]):
        set_cell_text(summary.cell(0, idx), header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="EDEFF3")

    summary_pairs = [
        ("数据表", str(len(report.tables)), "错误", str(len(errors))),
        ("警告", str(len(warnings)), "提取方式", extraction_label),
        ("原始字符", f"{diagnostics.get('raw_chars', 0):,}", "清洗后字符", f"{diagnostics.get('clean_chars', 0):,}"),
        ("压缩率", f"{diagnostics.get('compression_ratio', 0.0):.1%}", "疑似异常表", str(diagnostics.get("abnormal_table_count", 0))),
    ]
    for row_idx, values in enumerate(summary_pairs, start=1):
        for col_idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx % 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(summary.cell(row_idx, col_idx), value, align=align)

    if report.table_extraction_flags:
        diag_note = doc.add_paragraph()
        diag_run = diag_note.add_run(
            "提取提示："
            + "；".join(
                f"{report.tables[idx].monitoring_item}{f' ({report.tables[idx].borehole_id})' if report.tables[idx].borehole_id else ''}：{'；'.join(flags)}"
                for idx, flags in sorted(report.table_extraction_flags.items())
                if idx < len(report.tables)
            )
        )
        set_run_style(diag_run, size=10)
        style_paragraph(diag_note, size=10, space_before=4, space_after=6)

    doc.add_paragraph("")

    issues_heading = doc.add_paragraph()
    issues_run = issues_heading.add_run("三、问题清单")
    set_run_style(issues_run, size=13, bold=True)
    style_paragraph(issues_heading, size=13, bold=True, space_after=6)

    def add_issue_table(title_text: str, issue_list: list) -> None:
        title_p = doc.add_paragraph()
        title_r = title_p.add_run(title_text)
        set_run_style(title_r, size=11.5, bold=True)
        style_paragraph(title_p, size=11.5, bold=True, space_before=2, space_after=4)
        if not issue_list:
            empty_p = doc.add_paragraph()
            empty_r = empty_p.add_run("未发现需关注的问题。")
            set_run_style(empty_r, size=10.5)
            style_paragraph(empty_p, size=10.5, space_after=6)
            return
        table = doc.add_table(rows=1, cols=5, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["序号", "表名", "测点", "字段", "说明"]
        for idx, header in enumerate(headers):
            set_cell_text(table.cell(0, idx), header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="EDEFF3")
        for idx, issue in enumerate(issue_list, start=1):
            row = table.add_row().cells
            set_cell_text(row[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row[1], issue.table_name)
            set_cell_text(row[2], issue.point_id or "-", align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row[3], issue.field_name or "-", align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row[4], append_issue_source_hint(issue.message, issue.suspected_source))
        doc.add_paragraph("")

    add_issue_table("3.1 错误项", errors)
    add_issue_table("3.2 警告项", warnings)

    conclusion_heading = doc.add_paragraph()
    conclusion_run = conclusion_heading.add_run("四、结论")
    set_run_style(conclusion_run, size=13, bold=True)
    style_paragraph(conclusion_heading, size=13, bold=True, space_after=6)

    if report.conclusion:
        original_conclusion = doc.add_paragraph()
        original_run = original_conclusion.add_run(f"报告原文结论：{report.conclusion}")
        set_run_style(original_run, size=10.5)
        style_paragraph(original_conclusion, size=10.5, space_after=4)

    system_conclusion = doc.add_paragraph()
    if errors:
        result_text = f"自动检查结论：共发现 {len(errors)} 条错误、{len(warnings)} 条警告，建议结合原始报告进行人工复核。"
    elif warnings:
        result_text = f"自动检查结论：未发现错误，存在 {len(warnings)} 条警告，建议按需复核。"
    else:
        result_text = "自动检查结论：本次自动检查未发现错误或警告。"
    result_run = system_conclusion.add_run(result_text)
    set_run_style(result_run, size=10.5)
    style_paragraph(system_conclusion, size=10.5, space_after=10)

    footer = doc.add_paragraph()
    footer_run = footer.add_run("本文件由建筑变形监测报告核验台自动生成。")
    set_run_style(footer_run, size=9)
    style_paragraph(footer, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_html(md_content: str, project_name: str) -> str:
    """生成可打印的 HTML 报告"""
    import markdown

    html_body = markdown.markdown(
        md_content,
        extensions=["tables", "fenced_code"],
    )

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>{project_name} - 检查报告</title>
    <style>
        body {{ font-family: "Microsoft YaHei", "SimHei", sans-serif; max-width: 900px; margin: 0 auto; padding: 20px; color: #333; }}
        h1 {{ color: #1a5276; border-bottom: 3px solid #1a5276; padding-bottom: 10px; }}
        h2 {{ color: #2c3e50; border-bottom: 1px solid #bdc3c7; padding-bottom: 6px; margin-top: 30px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
        th {{ background-color: #f2f2f2; font-weight: bold; }}
        tr:nth-child(even) {{ background-color: #fafafa; }}
        blockquote {{ border-left: 4px solid #3498db; margin: 15px 0; padding: 10px 20px; background: #ecf6fd; }}
        code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-size: 0.9em; }}
        hr {{ border: none; border-top: 1px solid #ddd; margin: 30px 0; }}
        @media print {{
            body {{ max-width: 100%; padding: 0; }}
            h1 {{ page-break-before: avoid; }}
            table {{ page-break-inside: avoid; }}
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""


# ── 页面配置 ──────────────────────────────────────────
st.set_page_config(
    page_title="建筑变形监测报告核验台",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── 自定义样式 ────────────────────────────────────────
st.markdown("""
<style>
    :root {
        --panel-border: #d8dee9;
        --panel-bg: #ffffff;
        --muted: #5b6472;
        --heading: #162033;
        --accent: #1f5eff;
        --surface: #f5f7fb;
    }
    .stApp {
        background: linear-gradient(180deg, #f4f7fb 0%, #eef3f8 100%);
    }
    section[data-testid="stSidebar"] {
        background: #f7f9fc;
        border-right: 1px solid var(--panel-border);
    }
    .app-hero {
        background: linear-gradient(135deg, #ffffff 0%, #f4f8ff 100%);
        border: 1px solid var(--panel-border);
        border-radius: 18px;
        padding: 22px 24px;
        margin-bottom: 18px;
        box-shadow: 0 10px 28px rgba(15, 23, 42, 0.05);
    }
    .app-hero-kicker {
        font-size: 12px;
        letter-spacing: 0.08em;
        text-transform: uppercase;
        color: var(--accent);
        margin-bottom: 8px;
        font-weight: 600;
    }
    .app-hero h1 {
        margin: 0;
        color: var(--heading);
        font-size: 30px;
        line-height: 1.15;
        font-weight: 700;
    }
    .app-hero p {
        margin: 10px 0 0 0;
        color: var(--muted);
        font-size: 14px;
    }
    .phase-card {
        background: var(--panel-bg);
        border: 1px solid var(--panel-border);
        border-radius: 14px;
        padding: 14px 16px;
        margin: 10px 0 14px 0;
        box-shadow: 0 6px 18px rgba(15, 23, 42, 0.04);
    }
    .phase-step {
        font-size: 15px;
        font-weight: 600;
        color: var(--heading);
    }
    .phase-detail {
        margin-top: 6px;
        font-size: 13px;
        color: var(--muted);
    }
    .stMetric > div {
        border: 1px solid var(--panel-border);
        border-radius: 14px;
        padding: 14px;
        background: #ffffff;
        box-shadow: 0 6px 18px rgba(15, 23, 42, 0.04);
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        background: #edf2f7;
        border-radius: 10px;
        padding: 8px 14px;
        color: #334155;
    }
    .stTabs [aria-selected="true"] {
        background: #ffffff !important;
        border: 1px solid var(--panel-border);
        color: var(--heading) !important;
    }
    div[data-testid="stExpander"] details {
        border: 1px solid var(--panel-border);
        border-radius: 14px;
        background: #ffffff;
    }
    div[data-testid="stExpander"] details summary p {
        font-weight: 600;
        color: var(--heading);
    }
    .stButton button, .stDownloadButton button {
        border-radius: 12px;
        border: 1px solid var(--panel-border);
        min-height: 42px;
    }
    .welcome-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(260px, 1fr));
        gap: 16px;
        margin-top: 8px;
    }
    .info-card {
        background: #ffffff;
        border: 1px solid var(--panel-border);
        border-radius: 16px;
        padding: 18px 18px 16px 18px;
        box-shadow: 0 8px 24px rgba(15, 23, 42, 0.04);
        height: 100%;
    }
    .info-card h3 {
        margin: 0 0 10px 0;
        color: var(--heading);
        font-size: 16px;
    }
    .info-card p,
    .info-card li {
        color: var(--muted);
        font-size: 14px;
        line-height: 1.65;
    }
    .info-card ul,
    .info-card ol {
        margin: 0;
        padding-left: 18px;
    }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="app-hero">
  <div class="app-hero-kicker">Monitoring QA Console</div>
  <h1>建筑变形监测报告核验台</h1>
  <p>面向基坑监测报告的提取、结构化解析、计算校核、统计复核与审阅输出工作台。</p>
</div>
""", unsafe_allow_html=True)

# ── 侧边栏设置 ────────────────────────────────────────
with st.sidebar:
    st.header("运行设置")

    ocr_mode = st.radio(
        "PDF提取方式",
        ["优先 pdfplumber", "优先 PaddleOCR（表格）", "强制 PaddleOCR"],
        index=0,
        help="Web 默认先用 pdfplumber，文本质量不足时自动回退到 PaddleOCR；也可以显式优先或强制使用 OCR",
    )
    use_ocr = ocr_mode == "强制 PaddleOCR"
    prefer_ocr = ocr_mode == "优先 PaddleOCR（表格）" or use_ocr
    auto_fallback = not use_ocr

    st.divider()
    st.subheader("模型")
    from src.config import AVAILABLE_MODELS
    selected_model = st.selectbox(
        "结构化解析模型",
        AVAILABLE_MODELS,
        index=0,
        help="不同模型在表格理解、语义匹配和复核速度上各有差异",
    )

    st.divider()
    st.subheader("可选复核")
    do_self_verify = st.checkbox("LLM 复核错误项", value=False, help="按批次复核错误项，会显著增加等待时间")
    do_ai_review = st.checkbox("LLM 最终审核", value=False, help="对整份报告做最终审阅，耗时最长")
    st.caption("建议先查看确定性规则结果，必要时再开启复核。")

    st.divider()
    st.subheader("核心计算公式")
    st.code("本次变化 = 本次测值 − 上次测值\n累计变化 = 本次测值 − 初始测值\n变化速率 = 本次变化 / 时间(天)", language=None)
    st.caption("正负号表示方向，不表示绝对大小。")

# ── 文件上传区 ────────────────────────────────────────
uploaded = st.file_uploader(
    "上传监测报告 PDF",
    type=["pdf"],
    accept_multiple_files=False,
    help="支持文字版PDF和扫描件PDF",
)

if uploaded is not None:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(uploaded.read())
        tmp_path = tmp.name

    pdf_name = Path(uploaded.name).stem
    st.success(f"已载入文件：**{uploaded.name}** ({uploaded.size / 1024:.0f} KB)")

    if st.button("开始检查", type="primary", use_container_width=True):
        log_records.clear()

        # 运行时切换模型
        from src.config import set_model
        set_model(selected_model)

        # ── 实时进度容器 ──────────────────────────────
        progress_bar = st.progress(0)
        phase_placeholder = st.empty()
        status_container = st.status(f"正在处理（模型: {selected_model}）", expanded=True)
        start_time = time.time()

        def update_phase(step_label: str, detail: str = "", progress: int | None = None) -> None:
            detail_html = f"<div class='phase-detail'>{detail}</div>" if detail else ""
            phase_placeholder.markdown(
                f"<div class='phase-card'><div class='phase-step'>{step_label}</div>{detail_html}</div>",
                unsafe_allow_html=True,
            )
            if progress is not None:
                progress_bar.progress(progress)

        try:
            # ━━ Step 1: PDF 提取 ━━━━━━━━━━━━━━━━━━━━━
            update_phase("Step 1/8 · PDF 提取", "读取文本层，必要时回退 OCR。", 5)
            with status_container:
                st.write("Step 1/8 · 提取 PDF 内容")

            from src.tools.pdf_extractor import extract_pdf
            extraction_result = extract_pdf(
                tmp_path,
                use_ocr=use_ocr,
                prefer_ocr=prefer_ocr,
                auto_fallback=auto_fallback,
                ocr_output_dir=f"output/{pdf_name}_ocr_debug",
                return_details=True,
            )
            raw_text = extraction_result.text
            extraction_result.diagnostics.setdefault("method", extraction_result.method)
            extraction_result.diagnostics.setdefault("selected_profile", extraction_result.selected_profile)
            extraction_result.diagnostics.setdefault("debug_dir", extraction_result.debug_output_dir)

            update_phase(
                "Step 1/8 · PDF 提取",
                f"完成：{len(raw_text):,} 字符，方式 {extraction_result.method}/{extraction_result.selected_profile}",
                10,
            )
            with status_container:
                st.write(
                    "完成："
                    f"{len(raw_text):,} 字符 "
                    f"({extraction_result.method}/{extraction_result.selected_profile}, "
                    f"{extraction_result.diagnostics.get('raw_chars', 0):,} → "
                    f"{extraction_result.diagnostics.get('clean_chars', 0):,})"
                )

            # ━━ Step 2: LLM 结构化解析 ━━━━━━━━━━━━━━━
            update_phase("Step 2/8 · 结构化解析", "发送文本到 LLM，提取项目、阈值、汇总项和数据表。", 12)
            with status_container:
                st.write("Step 2/8 · 结构化解析")

            from src.tools.llm_parser import parse_report_with_llm
            report = parse_report_with_llm(raw_text)
            report.raw_text = raw_text
            report.extraction_diagnostics = extraction_result.diagnostics

            from src.tools.extraction_quality import analyze_extraction_quality
            analyze_extraction_quality(report)

            import src.config as cfg
            step_delay = getattr(cfg, "LLM_STEP_DELAY_SEC", 0)
            if step_delay > 0:
                with status_container:
                    st.write(f"等待 {step_delay} 秒以避免限流...")
                time.sleep(step_delay)

            from src.tools.table_analyzer import enrich_configs_with_llm
            enrich_configs_with_llm(report)

            update_phase(
                "Step 2/8 · 结构化解析",
                f"完成：{len(report.tables)} 张表，{len(report.thresholds)} 项阈值，{len(report.summary_items)} 项汇总。",
                35,
            )
            with status_container:
                st.write(f"完成：**{report.project_name}** — {len(report.tables)} 张表, "
                         f"{len(report.thresholds)} 项阈值, {len(report.summary_items)} 项汇总")

            # ━━ Step 2.5: 表格分析计划 ━━━━━━━━━━━━━━━━
            update_phase("Step 3/8 · 分析计划", "生成每张表的字段识别、单位判断和验证策略。", 38)
            with status_container:
                st.write("Step 3/8 · 分析表格结构与验证策略")

            from src.tools.table_analyzer import generate_analysis_plan
            analysis_plan = generate_analysis_plan(report)

            update_phase("Step 3/8 · 分析计划", f"完成：已生成 {len(analysis_plan)} 张表的验证策略。", 42)
            with status_container:
                st.write(f"完成：{len(analysis_plan)} 张表的验证策略已制定")

            # ━━ Step 4: 计算验证 ━━━━━━━━━━━━━━━━━━━━━
            update_phase("Step 4/8 · 计算验证", "逐条验证累计变化量、变化速率和深层位移变化。", 45)
            with status_container:
                st.write("Step 4/8 · 计算验证")

            from src.tools.calculation_checker import run_calculation_checks
            calc_issues = run_calculation_checks(report)

            update_phase("Step 4/8 · 计算验证", f"完成：发现 {len(calc_issues)} 个问题。", 55)
            with status_container:
                st.write(f"完成：{len(calc_issues)} 个问题")

            # ━━ Step 5: 统计验证 ━━━━━━━━━━━━━━━━━━━━━
            update_phase("Step 5/8 · 统计验证", "核对方向极值、速率极值和多页汇总引用。", 58)
            with status_container:
                st.write("Step 5/8 · 统计验证")

            from src.tools.statistics_checker import run_statistics_checks
            stats_issues = run_statistics_checks(report)

            update_phase("Step 5/8 · 统计验证", f"完成：发现 {len(stats_issues)} 个问题。", 63)
            with status_container:
                st.write(f"完成：{len(stats_issues)} 个问题")

            # ━━ Step 6: 逻辑检查 ━━━━━━━━━━━━━━━━━━━━━
            update_phase("Step 6/8 · 逻辑检查", "匹配阈值、汇总项与安全状态。", 66)
            with status_container:
                st.write("Step 6/8 · 逻辑检查")

            from src.tools.logic_checker import run_logic_checks
            logic_issues = run_logic_checks(report)

            update_phase("Step 6/8 · 逻辑检查", f"完成：发现 {len(logic_issues)} 个问题。", 70)
            with status_container:
                st.write(f"完成：{len(logic_issues)} 个问题")

            # ━━ Step 7: AI 自验证 ━━━━━━━━━━━━━━━━━━━━
            all_issues = calc_issues + stats_issues + logic_issues
            if do_self_verify:
                errors_to_verify = [i for i in all_issues if i.severity == "error"]
                if errors_to_verify:
                    step_delay = getattr(cfg, "LLM_STEP_DELAY_SEC", 0)
                    if step_delay > 0:
                        with status_container:
                            st.write(f"等待 {step_delay} 秒以避免限流...")
                        time.sleep(step_delay)

                    update_phase("Step 7/8 · 错误复核", f"待复核 {len(errors_to_verify)} 条错误。", 72)
                    with status_container:
                        st.write(f"Step 7/8 · 复核错误项（{len(errors_to_verify)} 条）")

                    def _on_self_verify_progress(event: dict) -> None:
                        stage = event.get("stage")
                        total_batches = max(event.get("total_batches", 1), 1)
                        batch_index = event.get("batch_index", 0)
                        if stage == "batch_start":
                            progress = 72 + int((batch_index - 1) / total_batches * 10)
                            update_phase(
                                "Step 7/8 · 错误复核",
                                f"正在处理第 {batch_index}/{total_batches} 批，本批 {event.get('batch_size', 0)} 条。",
                                progress,
                            )
                        elif stage == "batch_retry":
                            update_phase(
                                "Step 7/8 · 错误复核",
                                f"第 {batch_index}/{total_batches} 批重试中：{event.get('error', '')}",
                            )
                        elif stage == "batch_finish":
                            progress = 72 + int(batch_index / total_batches * 10)
                            update_phase(
                                "Step 7/8 · 错误复核",
                                f"已完成第 {batch_index}/{total_batches} 批，已降级 {event.get('downgraded', 0)}，已排除 {event.get('dismissed', 0)}。",
                                progress,
                            )
                        elif stage == "batch_failed":
                            update_phase(
                                "Step 7/8 · 错误复核",
                                f"第 {batch_index}/{total_batches} 批失败：{event.get('error', '')}",
                            )
                        elif stage == "done":
                            update_phase(
                                "Step 7/8 · 错误复核",
                                f"完成：共处理 {event.get('total_errors', 0)} 条，降级 {event.get('downgraded', 0)}，排除 {event.get('dismissed', 0)}。",
                                82,
                            )

                    from src.tools.self_verifier import verify_errors_with_llm
                    all_issues = verify_errors_with_llm(report, all_issues, progress_callback=_on_self_verify_progress)

                    with status_container:
                        st.write("完成：错误复核结束")
                else:
                    update_phase("Step 7/8 · 错误复核", "没有 error 级问题，跳过复核。", 82)
            else:
                update_phase("Step 7/8 · 错误复核", "已关闭该步骤。", 82)

            # ━━ Step 8: AI 最终审核 ━━━━━━━━━━━━━━━━━━
            ai_review = ""
            if do_ai_review:
                update_phase("Step 8/8 · 最终审核", "整理检查报告并发送最终审阅请求。", 84)
                with status_container:
                    st.write("Step 8/8 · 最终审核")

                from src.tools.report_generator import generate_report_md
                from src.tools.llm_parser import verify_report_with_llm
                prelim = generate_report_md(report, calc_issues, stats_issues, logic_issues, analysis_plan=analysis_plan)
                ai_review = verify_report_with_llm(
                    prelim,
                    raw_text,
                    progress_callback=lambda msg: update_phase("Step 8/8 · 最终审核", msg, 88),
                )

                update_phase("Step 8/8 · 最终审核", "完成：最终审核结果已返回。", 92)
                with status_container:
                    st.write("完成：最终审核结束")
            else:
                update_phase("Step 8/8 · 最终审核", "已关闭该步骤，直接生成报告。", 92)

            # ━━ 生成报告 ━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            update_phase("Step 8/8 · 生成报告", "整理最终报告与导出文件。", 94)
            from src.tools.report_generator import generate_report_md, save_report
            final_md = generate_report_md(report, calc_issues, stats_issues, logic_issues, ai_review, analysis_plan)
            output_path = f"output/{pdf_name}_检查报告.md"
            save_report(final_md, output_path)

            elapsed = time.time() - start_time
            progress_bar.progress(100)
            update_phase("处理完成", f"总耗时 {elapsed:.0f} 秒。", 100)
            status_container.update(label=f"检查完成，耗时 {elapsed:.0f} 秒", state="complete", expanded=False)

            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            # 结果展示区
            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

            errors = [i for i in all_issues if i.severity == "error"]
            warnings = [i for i in all_issues if i.severity == "warning"]
            infos = [i for i in all_issues if i.severity == "info"]

            st.markdown("---")
            st.subheader("检查结果总览")

            # ── 指标卡片 ──────────────────────────────
            c1, c2, c3, c4, c5 = st.columns(5)
            c1.metric("数据表", f"{len(report.tables)} 张")
            c2.metric("错误", f"{len(errors)}", delta=f"-{len(errors)}" if errors else None, delta_color="inverse")
            c3.metric("警告", f"{len(warnings)}")
            c4.metric("提示", f"{len(infos)}")
            c5.metric("耗时", f"{elapsed:.0f}s")
            _render_extraction_diagnostics(report)

            # ── 选项卡 ───────────────────────────────
            tab_report, tab_extract, tab_plan, tab_calc, tab_stats, tab_logic, tab_ai, tab_log = st.tabs([
                "检查报告", "提取诊断", "分析计划", "计算验证", "统计验证",
                "逻辑检查", "最终审核", "运行日志",
            ])

            with tab_report:
                st.markdown(final_md)

            with tab_extract:
                _render_extraction_diagnostics(report)

            with tab_plan:
                _render_analysis_plan(analysis_plan)

            with tab_calc:
                _render_issues("计算验证", calc_issues)

            with tab_stats:
                _render_issues("统计验证", stats_issues)

            with tab_logic:
                _render_issues("逻辑检查", logic_issues)

            with tab_ai:
                if ai_review:
                    st.markdown(ai_review)
                else:
                    st.info("当前未启用最终审核。")

            with tab_log:
                st.text("\n".join(log_records) if log_records else "暂无日志")

            # ── 导出按钮 ──────────────────────────────
            st.markdown("---")
            st.subheader("导出检查报告")

            dl_col1, dl_col2, dl_col3 = st.columns(3)

            with dl_col1:
                st.download_button(
                    label="下载 Markdown",
                    data=final_md,
                    file_name=f"{pdf_name}_检查报告.md",
                    mime="text/markdown",
                    use_container_width=True,
                )

            with dl_col2:
                docx_bytes = _generate_docx(final_md, report, errors, warnings)
                st.download_button(
                    label="下载 Word (docx)",
                    data=docx_bytes,
                    file_name=f"{pdf_name}_检查报告.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

            with dl_col3:
                html_content = _generate_html(final_md, report.project_name)
                st.download_button(
                    label="下载 HTML",
                    data=html_content,
                    file_name=f"{pdf_name}_检查报告.html",
                    mime="text/html",
                    use_container_width=True,
                )

        except Exception as e:
            progress_bar.empty()
            status_container.update(label="处理失败", state="error")
            st.error(f"处理过程中出错: {e}")
            logger.exception("处理失败")

            with st.expander("查看运行日志"):
                st.text("\n".join(log_records))

else:
    # ── 未上传文件时的欢迎页面 ────────────────────────
    col_left, col_right = st.columns([2, 1])

    with col_left:
        st.markdown("""
        <div class="info-card">
          <h3>开始使用</h3>
          <p>上传监测报告 PDF 后，系统将完成文本提取、结构化解析、规则校核与结果导出。</p>
        </div>
        """, unsafe_allow_html=True)
        st.markdown("""
        <div class="info-card">
          <h3>处理流程</h3>
          <ol>
            <li>文本提取：默认优先读取 PDF 文本层，必要时自动切换到 PaddleOCR。</li>
            <li>结构化解析：抽取项目、阈值、汇总项和监测数据表。</li>
            <li>分析计划：识别字段、单位、基准与验证口径。</li>
            <li>规则检查：执行计算验证、统计验证和逻辑检查。</li>
            <li>复核与导出：按需启用错误复核和最终审核，输出 Markdown、Word 与 HTML。</li>
          </ol>
        </div>
        """, unsafe_allow_html=True)

    with col_right:
        st.markdown("""
        <div class="welcome-grid">
          <div class="info-card">
            <h3>适用监测项</h3>
            <ul>
              <li>支护结构顶部水平位移、竖向位移</li>
              <li>周边地面沉降、道路沉降、管线沉降</li>
              <li>地下水位、锚索拉力、支撑轴力</li>
              <li>深层水平位移、立柱位移、裂缝监测</li>
            </ul>
          </div>
          <div class="info-card">
            <h3>运行说明</h3>
            <ul>
              <li>默认关闭两类远程复核步骤，优先返回确定性检查结果。</li>
              <li>第 7/8 步会显示批次、重试与完成状态，不再静默等待。</li>
              <li>导出的 Word 文档采用统一黑色字体与正式排版。</li>
            </ul>
          </div>
        </div>
        """, unsafe_allow_html=True)
